import streamlit as st
from docx import Document
from docx.shared import Pt
from docx.table import Table
from docx.text.paragraph import Paragraph
import io
import os

# --- 1. CONFIGURACIÓN DE PÁGINA ---
st.set_page_config(page_title="CRM Generator Pro", layout="wide")

# --- 2. DEFINICIÓN DE PLANTILLAS ---
TEMPLATES = {
    "M100 Minuta": "M100_CRM_Minuta v2 (2).docx",
    "M102 Gap Analysis": "M102_CRM_Gap_Analysis V2 (3).docx",
    "M101 Escenarios": "M101_CRM_Lista_de_escenarios_para_CRPUAT V2 (1).docx"
}

# --- 3. FUNCIONES DE APOYO ---
def aplicar_poppins(run, size=11):
    run.font.name = 'Poppins'
    run.font.size = Pt(size)

def iterar_bloques(parent):
    from docx.document import Document as _Document
    parent_elm = parent.element.body if isinstance(parent, _Document) else parent._tc
    for child in parent_elm.iterchildren():
        if child.tag.endswith('p'):
            yield Paragraph(child, parent)
        elif child.tag.endswith('tbl'):
            yield Table(child, parent)

# --- 4. EXTRACCIÓN MEJORADA (CORRECCIÓN AQUÍ) ---
def extraer_informacion(archivo_subido):
    datos = {k: "" for k in ["Fecha", "Objetivo", "Asistentes", "Puntos Discutidos", "Pendientes Cliente", "Pendientes Mycloud"]}
    if not archivo_subido: return datos
    
    try:
        doc = Document(archivo_subido)
        contexto = None
        
        for bloque in iterar_bloques(doc):
            if isinstance(bloque, Paragraph):
                txt = bloque.text.strip()
                txt_l = txt.lower()
                
                # Extracción de Fecha (Línea simple)
                if "fecha:" in txt_l:
                    datos["Fecha"] = txt.split(":", 1)[1].strip()
                
                # Detección de Secciones (Cambio de contexto)
                elif "objetivo:" in txt_l or "alcance:" in txt_l:
                    contexto = "Objetivo"
                    # Intentar extraer si el texto está en la misma línea
                    res = txt.split(":", 1)
                    if len(res) > 1 and res[1].strip():
                        datos["Objetivo"] = res[1].strip()
                elif "asistentes:" in txt_l:
                    contexto = "Asistentes"
                elif "puntos discutidos:" in txt_l:
                    contexto = "Puntos Discutidos"
                elif "pendientes del cliente" in txt_l or "pendientes cliente" in txt_l:
                    contexto = "Pendientes Cliente"
                elif "pendientes mycloud" in txt_l:
                    contexto = "Pendientes Mycloud"
                
                # Captura de contenido de párrafo (si no es un encabezado)
                elif txt and contexto:
                    # Si el contexto es Objetivo y ya capturamos algo en la misma línea, 
                    # evitamos duplicar si el párrafo es solo el encabezado
                    if contexto == "Objetivo" and (txt_l.startswith("objetivo") or txt_l.startswith("alcance")):
                        continue
                    datos[contexto] = (datos[contexto] + "\n" + txt).strip()
            
            elif isinstance(bloque, Table) and contexto:
                # Extraer filas de tablas (omitiendo encabezado)
                filas = [", ".join(c.text.strip() for c in r.cells if c.text.strip()) for r in bloque.rows[1:]]
                datos[contexto] = "\n".join(f for f in filas if f)
                
    except Exception as e:
        st.warning(f"Error al leer el documento: {e}")
    return datos

# --- 5. LÓGICA DE PROCESAMIENTO ---
def rellenar_tabla(tabla, texto_lineas, columnas):
    while len(tabla.rows) > 1:
        tabla._tbl.remove(tabla.rows[-1]._tr)
    for linea in texto_lineas.split('\n'):
        if not linea.strip(): continue
        nueva_fila = tabla.add_row().cells
        partes = linea.split(',')
        for i in range(min(len(partes), columnas)):
            nueva_fila[i].text = partes[i].strip()
            for p in nueva_fila[i].paragraphs:
                for run in p.runs: aplicar_poppins(run)

def procesar_word(template_path, datos, es_gap=False):
    doc = Document(template_path)
    for p in doc.paragraphs:
        if "Fecha:" in p.text:
            p.text = "Fecha: "
            aplicar_poppins(p.add_run(datos.get('Fecha', '')))
        elif "Objetivo:" in p.text or "Alcance:" in p.text:
            p.text = "Objetivo: " if not es_gap else "Objetivo : "
            aplicar_poppins(p.add_run(datos.get('Objetivo', '')))
        elif "Puntos discutidos:" in p.text and not es_gap:
            p.text = "Puntos discutidos:"
            for i, linea in enumerate(datos.get('Puntos Discutidos', '').split('\n'), 1):
                if linea.strip():
                    np = p.insert_paragraph_before(f"{i}. {linea.strip()}")
                    aplicar_poppins(np.runs[0] if np.runs else np.add_run())

    for tabla in doc.tables:
        h = tabla.cell(0,0).text.lower()
        if "nombre" in h and "puesto" in h: rellenar_tabla(tabla, datos.get("Asistentes", ""), 2)
        elif "pendientes del cliente" in h: rellenar_tabla(tabla, datos.get("Pendientes Cliente", ""), 3)
        elif "pendientes mycloud" in h: rellenar_tabla(tabla, datos.get("Pendientes Mycloud", ""), 3)
        elif "módulo" in h: rellenar_tabla(tabla, datos.get("Modulos", ""), 4)
        elif "entrega" in h or "pendientes" in h: rellenar_tabla(tabla, datos.get("Pendientes_Gap", ""), 3)
        elif "custom" in h: rellenar_tabla(tabla, datos.get("Custom", ""), 2)
        elif "web services" in h: rellenar_tabla(tabla, datos.get("WebServices", ""), 4)
        elif "workflows" in h: rellenar_tabla(tabla, datos.get("Workflows", ""), 5)
    return doc

# --- 6. INTERFAZ ---
with st.sidebar:
    st.header("🎨 Identidad Visual")
    logo_web = st.file_uploader("Sube tu logo:", type=["png", "jpg", "jpeg"])
    if logo_web: st.image(logo_web, use_container_width=True)
    elif os.path.exists("logo.png"): st.image("logo.png", use_container_width=True)
    st.divider()
    opcion = st.selectbox("Selecciona Plantilla:", list(TEMPLATES.keys()))

st.title("🚀 Generador CRM Profesional")

archivo_ref = st.file_uploader("📂 Sube la minuta anterior para auto-rellenar:", type=["docx"])
datos_auto = extraer_informacion(archivo_ref)

with st.form(key="main_form"):
    st.subheader(f"Edición de {opcion}")
    c1, c2 = st.columns(2)
    
    with c1:
        fecha = st.text_input("Fecha", value=datos_auto["Fecha"])
        asistentes = st.text_area("Asistentes (Nombre, Cargo)", value=datos_auto["Asistentes"], height=150)
        objetivo = st.text_area("Objetivo / Alcance", value=datos_auto["Objetivo"], height=100)

    with c2:
        if opcion == "M102 Gap Analysis":
            modulos = st.text_area("Módulos (Item, Nombre, Desc, Estatus)")
            pend_gap = st.text_area("Pendientes/Entrega (Tarea, Resp, Fecha)")
            custom = st.text_area("Custom Functions (Item, Desc)")
            ws = st.text_area("Web Services (Item, Nombre, Tipo, Param)")
            wf = st.text_area("Workflows (Item, Módulo, Cuándo, Qué, Acciones)")
        else:
            puntos = st.text_area("Puntos Discutidos", value=datos_auto["Puntos Discutidos"], height=150)
            p_cli = st.text_area("Pendientes Cliente", value=datos_auto["Pendientes Cliente"], height=100)
            p_my = st.text_area("Pendientes Mycloud", value=datos_auto["Pendientes Mycloud"], height=100)

    btn = st.form_submit_button("🔨 GENERAR DOCUMENTO")

if btn:
    es_gap = (opcion == "M102 Gap Analysis")
    final_dict = {
        "Fecha": fecha, "Asistentes": asistentes, "Objetivo": objetivo,
        "Puntos Discutidos": puntos if not es_gap else "",
        "Pendientes Cliente": p_cli if not es_gap else "",
        "Pendientes Mycloud": p_my if not es_gap else "",
        "Modulos": modulos if es_gap else "",
        "Pendientes_Gap": pend_gap if es_gap else "",
        "Custom": custom if es_gap else "",
        "WebServices": ws if es_gap else "",
        "Workflows": wf if es_gap else ""
    }

    try:
        resultado = procesar_word(TEMPLATES[opcion], final_dict, es_gap=es_gap)
        buf = io.BytesIO()
        resultado.save(buf)
        buf.seek(0)
        st.success("✅ ¡Documento generado!")
        st.download_button("📥 Descargar Archivo Word", buf, f"{opcion}.docx")
    except Exception as e:
        st.error(f"Error al generar: {e}")
